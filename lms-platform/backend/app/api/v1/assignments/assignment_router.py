"""
Assignment Router
Handles assignment management and submissions
"""

import os, uuid, io
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, status, Query, Form, UploadFile, File
from fastapi.responses import FileResponse, Response
from sqlalchemy.orm import Session

from ....core.database import get_db_session
from ....core.security import get_current_user
from ....core.dependencies import require_teacher, require_student, require_teacher_or_student
from ....core.exceptions import NotFoundError, ConflictError, ValidationError
from ....schemas.assignment import AssignmentCreate, AssignmentUpdate, AssignmentResponse
from ....schemas.submission import SubmissionCreate, SubmissionUpdate, SubmissionResponse
from ....schemas.common import ResponseWrapper
from ....services.assignment_service import AssignmentService
from ....services.submission_service import SubmissionService
from ....services.notification_service import NotificationService
from ....models import User, Assignment, AssignmentSubmission, CourseEnrollment, Course
from ....core.config import settings

router = APIRouter()


# ==================== Assignment Routes ====================

@router.get("/", response_model=ResponseWrapper)
async def get_assignments(
    course_id: Optional[int] = Query(None, description="Filter by course"),
    is_published: Optional[bool] = Query(None, description="Filter by published status"),
    page: int = Query(1, ge=1, description="Page number"),
    page_size: int = Query(20, ge=1, le=100, description="Items per page"),
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher_or_student)
):
    """
    Get assignments with filters and pagination
    """
    skip = (page - 1) * page_size
    
    if course_id:
        assignments, total = AssignmentService.get_course_assignments(
            db,
            course_id,
            skip=skip,
            limit=page_size,
            is_published=is_published
        )
    else:
        # Get all assignments with filters
        query = db.query(Assignment).filter(Assignment.deleted_at.is_(None))
        if is_published is not None:
            query = query.filter(Assignment.is_published == is_published)
        
        # If student, only show assignments from their courses
        if current_user.get("role") == "student":
            from ....models import CourseEnrollment
            enrolled_courses = db.query(CourseEnrollment.course_id).filter(
                CourseEnrollment.student_id == current_user["user_id"],
                CourseEnrollment.status == "active"
            ).subquery()
            query = query.filter(Assignment.course_id.in_(enrolled_courses))
        
        total = query.count()
        assignments = query.order_by(
            Assignment.deadline.asc()
        ).offset(skip).limit(page_size).all()
    
    # Convert to response
    assignment_responses = []
    for assignment in assignments:
        # ✅ Yeh change karein - AssignmentSubmission use karein
        submissions_count = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id == assignment.id,
            AssignmentSubmission.deleted_at.is_(None)
        ).count()
        
        graded_count = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id == assignment.id,
            AssignmentSubmission.status == "graded",
            AssignmentSubmission.deleted_at.is_(None)
        ).count()
        
        assignment_responses.append(AssignmentResponse(
            id=assignment.id,
            course_id=assignment.course_id,
            course_title=assignment.course.title if assignment.course else None,
            title=assignment.title,
            description=assignment.description,
            instructions=assignment.instructions,
            max_score=assignment.max_score,
            deadline=assignment.deadline,
            is_published=assignment.is_published,
            allow_late_submission=assignment.allow_late_submission if assignment.allow_late_submission is not None else False,
            late_submission_penalty=assignment.late_submission_penalty if assignment.late_submission_penalty is not None else 0.0,
            max_file_size=assignment.max_file_size if assignment.max_file_size is not None else 10485760,
            allowed_file_types=assignment.allowed_file_types if assignment.allowed_file_types is not None else ".pdf,.doc,.docx,.zip",
            created_at=assignment.created_at,
            updated_at=assignment.updated_at,
            submissions_count=submissions_count,
            graded_count=graded_count
        ))
    
    total_pages = (total + page_size - 1) // page_size
    
    return ResponseWrapper(
        success=True,
        message="Assignments retrieved successfully",
        data={
            "assignments": assignment_responses,
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": total_pages
        }
    )


@router.get("/my", response_model=ResponseWrapper)
async def get_my_assignments(
    status_filter: Optional[str] = Query(None, alias="status", description="Filter: pending, submitted, graded"),
    page: int = Query(1, ge=1, description="Page number"),
    page_size: int = Query(50, ge=1, le=100, description="Items per page"),
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_student)
):
    """
    Get assignments for the current student with their submission status.
    - pending: published assignments from enrolled courses with NO submission
    - submitted: assignments the student has submitted (not yet graded)
    - graded: assignments the student has submitted and been graded
    - None: return all
    """
    skip = (page - 1) * page_size

    enrolled_course_ids = [
        row[0] for row in db.query(CourseEnrollment.course_id).filter(
            CourseEnrollment.student_id == current_user["user_id"],
            CourseEnrollment.status == "active"
        ).all()
    ]

    if not enrolled_course_ids:
        return ResponseWrapper(
            success=True,
            message="No enrolled courses",
            data={"assignments": [], "total": 0, "page": page, "page_size": page_size, "total_pages": 0}
        )

    assignments = db.query(Assignment).filter(
        Assignment.course_id.in_(enrolled_course_ids),
        Assignment.is_published == True,
        Assignment.deleted_at.is_(None)
    ).order_by(Assignment.deadline.is_(None), Assignment.deadline.asc()).all()

    student_id = current_user["user_id"]
    result = []
    for a in assignments:
        submission = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id == a.id,
            AssignmentSubmission.student_id == student_id,
            AssignmentSubmission.deleted_at.is_(None)
        ).first()

        if submission:
            my_status = "graded" if submission.status == "graded" else "submitted"
        else:
            is_overdue = a.deadline and datetime.now() > a.deadline
            my_status = "overdue" if is_overdue else "pending"

        if status_filter and my_status != status_filter:
            continue

        result.append({
            "id": a.id,
            "title": a.title,
            "description": a.description,
            "course_id": a.course_id,
            "course_title": a.course.title if a.course else None,
            "deadline": a.deadline.isoformat() if a.deadline else None,
            "max_score": a.max_score,
            "assignment_type": getattr(a, 'assignment_type', 'general'),
            "instructions": a.instructions,
            "created_at": a.created_at.isoformat() if a.created_at else None,
            "my_status": my_status,
            "submission_id": submission.id if submission else None,
            "grade": float(submission.grade) if submission and submission.grade is not None else None,
            "feedback": submission.feedback if submission else None,
            "submitted_at": submission.submitted_at.isoformat() if submission and submission.submitted_at else None,
        })

    total = len(result)
    paginated = result[skip:skip + page_size]
    total_pages = (total + page_size - 1) // page_size

    return ResponseWrapper(
        success=True,
        message="My assignments retrieved successfully",
        data={
            "assignments": paginated,
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": total_pages
        }
    )


# ==================== Baqi endpoints same rahenge ====================

@router.get("/{assignment_id}", response_model=ResponseWrapper)
async def get_assignment(
    assignment_id: int,
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher_or_student)
):
    """
    Get a single assignment by ID
    """
    assignment = db.query(Assignment).filter(
        Assignment.id == assignment_id,
        Assignment.deleted_at.is_(None)
    ).first()

    if not assignment:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "NOT_FOUND", "message": "Assignment not found"}}
        )

    if current_user.get("role") == "student":
        enrollment = db.query(CourseEnrollment).filter(
            CourseEnrollment.student_id == current_user["user_id"],
            CourseEnrollment.course_id == assignment.course_id,
            CourseEnrollment.status == "active"
        ).first()
        if not enrollment:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail={"error": {"code": "FORBIDDEN", "message": "Not enrolled in this course"}}
            )

    return ResponseWrapper(
        success=True,
        message="Assignment retrieved successfully",
        data=AssignmentResponse(
            id=assignment.id,
            course_id=assignment.course_id,
            course_title=assignment.course.title if assignment.course else None,
            title=assignment.title,
            description=assignment.description,
            instructions=assignment.instructions,
            max_score=assignment.max_score,
            deadline=assignment.deadline,
            is_published=assignment.is_published,
            allow_late_submission=assignment.allow_late_submission if assignment.allow_late_submission is not None else False,
            late_submission_penalty=assignment.late_submission_penalty if assignment.late_submission_penalty is not None else 0.0,
            max_file_size=assignment.max_file_size if assignment.max_file_size is not None else 10485760,
            allowed_file_types=assignment.allowed_file_types if assignment.allowed_file_types is not None else ".pdf,.doc,.docx,.zip",
            created_at=assignment.created_at,
            updated_at=assignment.updated_at
        )
    )


@router.post("/", response_model=ResponseWrapper)
async def create_assignment(
    assignment_data: AssignmentCreate,
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher)
):
    """
    Create a new assignment (Teacher/Admin only)
    """
    try:
        assignment = AssignmentService.create_assignment(db, assignment_data)

        course = db.query(Course).filter(Course.id == assignment.course_id).first()
        if course:
            enrolled = db.query(CourseEnrollment).filter(
                CourseEnrollment.course_id == assignment.course_id,
                CourseEnrollment.status == 'active'
            ).all()
            for e in enrolled:
                NotificationService.create_notification(
                    db,
                    user_id=e.student_id,
                    title="New Assignment Posted",
                    message=f"New assignment '{assignment.title}' in {course.title} — due soon!",
                    type="info",
                    link=f"/student/assignments/{assignment.id}"
                )

        return ResponseWrapper(
            success=True,
            message="Assignment created successfully",
            data=AssignmentResponse(
                id=assignment.id,
                course_id=assignment.course_id,
                course_title=assignment.course.title if assignment.course else None,
                title=assignment.title,
                description=assignment.description,
                instructions=assignment.instructions,
                max_score=assignment.max_score,
                deadline=assignment.deadline,
                is_published=assignment.is_published,
                allow_late_submission=assignment.allow_late_submission,
                late_submission_penalty=assignment.late_submission_penalty,
                max_file_size=assignment.max_file_size,
                allowed_file_types=assignment.allowed_file_types,
                created_at=assignment.created_at,
                updated_at=assignment.updated_at
            )
        )
    except NotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "NOT_FOUND", "message": str(e)}}
        )


@router.put("/{assignment_id}", response_model=ResponseWrapper)
async def update_assignment(
    assignment_id: int,
    assignment_data: AssignmentUpdate,
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher)
):
    """
    Update an assignment (Teacher/Admin only)
    """
    try:
        existing_assignment = AssignmentService.get_assignment_by_id(db, assignment_id)
        if not existing_assignment:
            raise NotFoundError("Assignment not found", resource_type="Assignment", resource_id=assignment_id)
        if current_user.get("role") == "teacher":
            course = db.query(Course).filter(Course.id == existing_assignment.course_id).first()
            if not course or course.teacher_id != current_user["user_id"]:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail={"error": {"code": "FORBIDDEN", "message": "You can only update assignments in your own courses"}}
                )

        assignment = AssignmentService.update_assignment(db, assignment_id, assignment_data)
        
        return ResponseWrapper(
            success=True,
            message="Assignment updated successfully",
            data=AssignmentResponse(
                id=assignment.id,
                course_id=assignment.course_id,
                course_title=assignment.course.title if assignment.course else None,
                title=assignment.title,
                description=assignment.description,
                instructions=assignment.instructions,
                max_score=assignment.max_score,
                deadline=assignment.deadline,
                is_published=assignment.is_published,
                allow_late_submission=assignment.allow_late_submission,
                late_submission_penalty=assignment.late_submission_penalty,
                max_file_size=assignment.max_file_size,
                allowed_file_types=assignment.allowed_file_types,
                created_at=assignment.created_at,
                updated_at=assignment.updated_at
            )
        )
    except NotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "NOT_FOUND", "message": str(e)}}
        )


@router.delete("/{assignment_id}", response_model=ResponseWrapper)
async def delete_assignment(
    assignment_id: int,
    permanent: bool = Query(False, description="Permanently delete"),
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher)
):
    """
    Delete an assignment (Teacher/Admin only)
    """
    try:
        existing_assignment = AssignmentService.get_assignment_by_id(db, assignment_id)
        if not existing_assignment:
            raise NotFoundError("Assignment not found", resource_type="Assignment", resource_id=assignment_id)
        if current_user.get("role") == "teacher":
            course = db.query(Course).filter(Course.id == existing_assignment.course_id).first()
            if not course or course.teacher_id != current_user["user_id"]:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail={"error": {"code": "FORBIDDEN", "message": "You can only delete assignments in your own courses"}}
                )

        AssignmentService.delete_assignment(db, assignment_id, soft_delete=not permanent)
        
        return ResponseWrapper(
            success=True,
            message="Assignment deleted successfully"
        )
    except NotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "NOT_FOUND", "message": str(e)}}
        )


# ==================== Submission Routes ====================

@router.get("/{assignment_id}/submissions/my", response_model=ResponseWrapper)
async def get_my_submission(
    assignment_id: int,
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_student)
):
    """
    Get current student's submission for a specific assignment
    """
    submission = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.assignment_id == assignment_id,
        AssignmentSubmission.student_id == current_user["user_id"],
        AssignmentSubmission.deleted_at.is_(None)
    ).first()

    if not submission:
        return ResponseWrapper(
            success=True,
            message="No submission found",
            data=None
        )

    download_url = f"/api/v1/assignments/submissions/{submission.id}/download" if submission.file_path else ""
    return ResponseWrapper(
        success=True,
        message="Submission retrieved successfully",
        data=SubmissionResponse(
            id=submission.id,
            assignment_id=submission.assignment_id,
            assignment_title=submission.assignment.title if submission.assignment else None,
            student_id=submission.student_id,
            student_name=submission.student.profile.full_name if submission.student and submission.student.profile else None,
            file_name=submission.file_name,
            file_path=download_url,
            file_size=submission.file_size,
            mime_type=submission.mime_type,
            submission_text=submission.submission_text,
            is_late=submission.is_late,
            plagiarism_score=submission.plagiarism_score,
            status=submission.status,
            submitted_at=submission.submitted_at,
            graded_at=submission.graded_at,
            created_at=submission.created_at,
            updated_at=submission.updated_at,
            grade=float(submission.grade) if submission.grade is not None else None,
            feedback=submission.feedback
        )
    )


@router.get("/{assignment_id}/submissions", response_model=ResponseWrapper)
async def get_submissions(
    assignment_id: int,
    status_filter: Optional[str] = Query(None, alias="status", description="Filter by status"),
    page: int = Query(1, ge=1, description="Page number"),
    page_size: int = Query(20, ge=1, le=100, description="Items per page"),
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher)
):
    """
    Get submissions for an assignment (Teacher/Admin only)
    """
    try:
        skip = (page - 1) * page_size
        
        submissions, total = SubmissionService.get_assignment_submissions(
            db,
            assignment_id,
            skip=skip,
            limit=page_size,
            status=status_filter
        )
        
        submission_responses = []
        for submission in submissions:
            submission_responses.append(SubmissionResponse(
                id=submission.id,
                assignment_id=submission.assignment_id,
                assignment_title=submission.assignment.title if submission.assignment else None,
                student_id=submission.student_id,
                student_name=submission.student.profile.full_name if submission.student and submission.student.profile else None,
                file_name=submission.file_name,
                file_path=f"/api/v1/assignments/submissions/{submission.id}/download" if submission.file_path else "",
                file_size=submission.file_size,
                mime_type=submission.mime_type,
                submission_text=submission.submission_text,
                is_late=submission.is_late,
                plagiarism_score=submission.plagiarism_score,
                status=submission.status,
                submitted_at=submission.submitted_at,
                graded_at=submission.graded_at,
                created_at=submission.created_at,
                updated_at=submission.updated_at,
                grade=submission.grade,
                feedback=submission.feedback
            ))
        
        total_pages = (total + page_size - 1) // page_size
        
        return ResponseWrapper(
            success=True,
            message="Submissions retrieved successfully",
            data={
                "submissions": submission_responses,
                "total": total,
                "page": page,
                "page_size": page_size,
                "total_pages": total_pages
            }
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": {"code": "ERROR", "message": str(e)}}
        )


@router.post("/{assignment_id}/submit", response_model=ResponseWrapper)
async def submit_assignment(
    assignment_id: int,
    content: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_student)
):
    """
    Submit an assignment (Student only)
    Accepts multipart/form-data with optional 'content' (text) and 'file' (upload)
    """
    try:
        file_name = ""
        file_path = ""
        file_size = 0
        mime_type = ""

        if file and file.filename:
            upload_dir = os.path.join(settings.UPLOAD_DIR, "submissions")
            os.makedirs(upload_dir, exist_ok=True)
            ext = os.path.splitext(file.filename)[1] or ""
            safe_name = f"{uuid.uuid4().hex}{ext}"
            dest = os.path.join(upload_dir, safe_name)
            file_bytes = await file.read()
            with open(dest, "wb") as f:
                f.write(file_bytes)
            file_name = file.filename
            file_path = dest
            file_size = len(file_bytes)
            mime_type = file.content_type or "application/octet-stream"

        sub_data = SubmissionCreate(
            assignment_id=assignment_id,
            student_id=current_user["user_id"],
            file_name=file_name,
            file_path=file_path,
            file_size=file_size,
            mime_type=mime_type,
            submission_text=content
        )

        submission = SubmissionService.submit_assignment(db, sub_data)

        try:
            assignment_rec = None
            if submission.assignment:
                assignment_rec = submission.assignment
            elif assignment_id:
                assignment_rec = db.query(Assignment).filter(Assignment.id == assignment_id).first()

            if assignment_rec and assignment_rec.course_id:
                course = db.query(Course).filter(Course.id == assignment_rec.course_id).first()
                student_name = submission.student.profile.full_name if submission.student and submission.student.profile else "A student"
                if course and course.teacher_id:
                    NotificationService.create_notification(
                        db,
                        user_id=course.teacher_id,
                        title="New Submission Received",
                        message=f"{student_name} submitted '{assignment_rec.title}'",
                        type="info",
                        link=f"/teacher/assignments/{assignment_rec.id}/submissions"
                    )
                admins = db.query(User).filter(User.role == "admin", User.deleted_at.is_(None)).all()
                for admin in admins:
                    NotificationService.create_notification(
                        db,
                        user_id=admin.id,
                        title="New Submission Received",
                        message=f"{student_name} submitted '{assignment_rec.title}'",
                        type="info",
                        link=f"/teacher/assignments/{assignment_rec.id}/submissions"
                    )
        except Exception:
            pass

        return ResponseWrapper(
            success=True,
            message="Assignment submitted successfully",
            data=SubmissionResponse(
                id=submission.id,
                assignment_id=submission.assignment_id,
                assignment_title=submission.assignment.title if submission.assignment else None,
                student_id=submission.student_id,
                student_name=submission.student.profile.full_name if submission.student and submission.student.profile else None,
                file_name=submission.file_name,
                file_path=f"/api/v1/assignments/submissions/{submission.id}/download" if submission.file_path else "",
                file_size=submission.file_size,
                mime_type=submission.mime_type,
                submission_text=submission.submission_text,
                is_late=submission.is_late,
                plagiarism_score=submission.plagiarism_score,
                status=submission.status,
                submitted_at=submission.submitted_at,
                graded_at=submission.graded_at,
                created_at=submission.created_at,
                updated_at=submission.updated_at
            )
        )
    except NotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=str(e)
        )
    except ConflictError as e:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=str(e)
        )
    except ValidationError as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=str(e)
        )


@router.put("/submissions/{submission_id}", response_model=ResponseWrapper)
async def update_submission(
    submission_id: int,
    submission_data: SubmissionUpdate,
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_student)
):
    """
    Update a submission (Student only - before grading)
    """
    try:
        submission = SubmissionService.update_submission(db, submission_id, submission_data)
        
        return ResponseWrapper(
            success=True,
            message="Submission updated successfully",
            data=SubmissionResponse(
                id=submission.id,
                assignment_id=submission.assignment_id,
                assignment_title=submission.assignment.title if submission.assignment else None,
                student_id=submission.student_id,
                student_name=submission.student.profile.full_name if submission.student and submission.student.profile else None,
                file_name=submission.file_name,
                file_path=submission.file_path,
                file_size=submission.file_size,
                mime_type=submission.mime_type,
                submission_text=submission.submission_text,
                is_late=submission.is_late,
                plagiarism_score=submission.plagiarism_score,
                status=submission.status,
                submitted_at=submission.submitted_at,
                graded_at=submission.graded_at,
                created_at=submission.created_at,
                updated_at=submission.updated_at
            )
        )
    except NotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "NOT_FOUND", "message": str(e)}}
        )
    except ValidationError as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"error": {"code": "VALIDATION_ERROR", "message": str(e)}}
        )


@router.get("/submissions/{submission_id}/download")
async def download_submission_file(
    submission_id: int,
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher_or_student)
):
    """Download the submitted file (Teacher/Student/Admin)"""
    submission = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.id == submission_id,
        AssignmentSubmission.deleted_at.is_(None)
    ).first()
    if not submission or not submission.file_path or not os.path.exists(submission.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    if current_user.get("role") == "student" and submission.student_id != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Access denied")
    return FileResponse(submission.file_path, filename=submission.file_name or "download")


@router.get("/submissions/{submission_id}/export")
async def export_submission_text(
    submission_id: int,
    format: str = Query("docx", description="Export format: docx, html"),
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher_or_student)
):
    """Export text-based submission as .docx or .html (Teacher/Student/Admin)"""
    submission = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.id == submission_id,
        AssignmentSubmission.deleted_at.is_(None)
    ).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    if current_user.get("role") == "student" and submission.student_id != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Access denied")
    text = submission.submission_text or ""
    student_name = submission.student.profile.full_name if submission.student and submission.student.profile else "Student"
    title = submission.assignment.title if submission.assignment else "Submission"

    if format == "html":
        html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><title>{title}</title></head>
<body><h1>{title}</h1><p><strong>Student:</strong> {student_name}</p><hr>{text}</body></html>"""
        return Response(content=html, media_type="text/html", headers={"Content-Disposition": f"attachment; filename=\"{title}.html\""})

    # docx
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Student: {student_name}")
        doc.add_paragraph("─" * 40)
        if text:
            from docx.oxml.ns import qn
            body = doc.add_paragraph()
            run = body.add_run(text)
            run.font.size = Pt(11)
        else:
            doc.add_paragraph("(No text content)")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return Response(content=buf.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=\"{title}.docx\""})
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX export library not available")


from pydantic import BaseModel as _BaseModel


class _GradeBody(_BaseModel):
    score: float
    feedback: Optional[str] = None


@router.post("/submissions/{submission_id}/grade", response_model=ResponseWrapper)
async def grade_submission(
    submission_id: int,
    body: _GradeBody,
    db: Session = Depends(get_db_session),
    current_user: dict = Depends(require_teacher)
):
    """
    Grade a submission (Teacher/Admin only)
    """
    try:
        submission = SubmissionService.grade_submission(db, submission_id, body.score, body.feedback)

        NotificationService.create_notification(
            db,
            user_id=submission.student_id,
            title="Assignment Graded",
            message=f"Your assignment '{submission.assignment.title if submission.assignment else ''}' has been graded — Score: {body.score}",
            type="success",
            link=f"/student/assignments/{submission.assignment_id}"
        )

        return ResponseWrapper(
            success=True,
            message="Submission graded successfully",
            data=SubmissionResponse(
                id=submission.id,
                assignment_id=submission.assignment_id,
                assignment_title=submission.assignment.title if submission.assignment else None,
                student_id=submission.student_id,
                student_name=submission.student.profile.full_name if submission.student and submission.student.profile else None,
                file_name=submission.file_name,
                file_path=submission.file_path,
                file_size=submission.file_size,
                mime_type=submission.mime_type,
                submission_text=submission.submission_text,
                is_late=submission.is_late,
                plagiarism_score=submission.plagiarism_score,
                status=submission.status,
                submitted_at=submission.submitted_at,
                graded_at=submission.graded_at,
                created_at=submission.created_at,
                updated_at=submission.updated_at,
                grade=float(submission.grade) if submission.grade is not None else None,
                feedback=submission.feedback
            )
        )
    except NotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "NOT_FOUND", "message": str(e)}}
        )
    except ValidationError as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"error": {"code": "VALIDATION_ERROR", "message": str(e)}}
        )